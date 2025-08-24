
import openai
import llama_index
from llama_index.llms.openai import OpenAI

try:
    from llama_index import (
        VectorStoreIndex,
        ServiceContext,
        Document,
        SimpleDirectoryReader,
    )
except ImportError:
    from llama_index.core import (
        VectorStoreIndex,
        ServiceContext,
        Document,
        SimpleDirectoryReader,
    )

from azure.storage.blob import BlobServiceClient
from io import BytesIO

from llama_index.core.extractors import (
    TitleExtractor,
    QuestionsAnsweredExtractor,
)
from llama_index.core.node_parser import TokenTextSplitter
import streamlit as st
import logging
import os 
import dotenv
from dotenv import load_dotenv

load_dotenv()

azure_storage_account_name = os.environ["AZURE_STORAGE_ACCOUNT_NAME"]
azure_storage_account_key = os.environ["AZURE_STORAGE_ACCOUNT_KEY"]
connection_string_blob = os.environ["CONNECTION_STRING_BLOB"]
container_name = None
logging_container_name = os.environ["LOGGING_CONTAINER_NAME"]
blob_service_client = BlobServiceClient.from_connection_string(f"DefaultEndpointsProtocol=https;AccountName={azure_storage_account_name};AccountKey={azure_storage_account_key}")



class AzureBlobStorageHandler(logging.Handler):
    def __init__(self, connection_string, container_name, blob_name):
        super().__init__()
        self.connection_string = connection_string
        self.container_name = container_name
        self.blob_name = blob_name
        self.blob_service_client = blob_service_client
        self.container_client = self.blob_service_client.get_container_client(
            container_name
        )

    def emit(self, record):
        log_entry = self.format(record)
        self.append_log_to_blob(log_entry)

    def append_log_to_blob(self, log_entry):
        blob_client = self.container_client.get_blob_client(self.blob_name)
        try:
            blob_data = blob_client.download_blob().content_as_text()
        except:
            blob_data = ""
        updated_log_data = blob_data + "\n" + log_entry
        blob_client.upload_blob(updated_log_data, overwrite=True)


class Singleton(type):
    _instances = {}
    def __call__(cls, *args, **kwargs):
        if cls not in cls._instances:
            cls._instances[cls] = super(Singleton, cls).__call__(*args, **kwargs)
        return cls._instances[cls]

class Logger(metaclass=Singleton):
    def __init__(self):
        self.logger = logging.getLogger("azureLogger")
        self.logger.setLevel(logging.DEBUG)
        azure_handler = AzureBlobStorageHandler(
            connection_string_blob, logging_container_name, "test.log"
        )
        azure_handler.setLevel(logging.DEBUG)
        formatter = logging.Formatter(
            "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
        )
        azure_handler.setFormatter(formatter)
        self.logger.addHandler(azure_handler)

    def get_logger(self):
        return self.logger


def list_all_containers():
    container_list = list()
    containers = blob_service_client.list_containers()
    for container in containers:
        if "genai" in container.name:
            container_list.append(container.name)
    return container_list



def list_all_files(container_name):
    container = blob_service_client.get_container_client(container_name)
    blobs = container.list_blobs()
    blob_list_display = []
    for blob in blobs:
        blob_list_display.append({
            "Name": blob.name,
            "Size": blob.size        
        })
    return blob_list_display


def upload_to_azure_storage(file,container_name):
    blob_client = blob_service_client.get_blob_client(
        container=container_name, blob=file.name
    )
    blob_client.upload_blob(file)
    return True


def delete_all_files(container_name):
    container_client = blob_service_client.get_container_client(container_name)
    blob_list = container_client.list_blobs()
    for blob in blob_list:
        container_client.delete_blob(blob.name)
    return True


def create_new_container(container_name):
    # genai_container = f"genai-{container_name}"
    genai_container = container_name
    blob_service_client.create_container(genai_container)
    return True




import tempfile
import time
import math
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from llama_index.core.readers.base import BaseReader
from llama_index.core.schema import Document

class SimplePDFReader(BaseReader):
    """Custom PDF Parser that keeps full PDFs as a single Document"""
    
    def load_data(
        self,
        file: Path,
        extra_info: Optional[Dict] = None
    ) -> List[Document]:
        """Parse file and return as single document."""
        try:
            import pypdf
        except ImportError:
            raise ImportError(
                "pypdf is required to read PDF files: `pip install pypdf`"
            )
        
        with open(file, "rb") as fp:
            # Create a PDF object
            pdf = pypdf.PdfReader(fp)
            # Get the number of pages in the PDF document
            num_pages = len(pdf.pages)
            
            # Concatenate all pages into single text
            full_text = ""
            for page in range(num_pages):
                page_text = pdf.pages[page].extract_text()
                full_text += page_text + "\n\n"  # Add spacing between pages
            
            # Create metadata
            metadata = {"file_name": str(file.name), "num_pages": num_pages}
            if extra_info:
                metadata.update(extra_info)
            
            # Return single document with full text
            return [Document(text=full_text.strip(), metadata=metadata)]

class SimpleDocxReader(BaseReader):
    """Custom DOCX Parser that keeps full documents as single Document"""
    
    def load_data(
        self,
        file: Path,
        extra_info: Optional[Dict] = None
    ) -> List[Document]:
        """Parse DOCX file and return as single document."""
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is required to read DOCX files: `pip install python-docx`"
            )
        
        doc = docx.Document(file)
        full_text = ""
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Create metadata
        metadata = {"file_name": str(file.name)}
        if extra_info:
            metadata.update(extra_info)
        
        # Return single document with full text
        return [Document(text=full_text.strip(), metadata=metadata)]

class CustomAzStorageBlobReader(BaseReader):
    """Azure Storage Blob reader that loads full documents without chunking."""
    
    def __init__(
        self,
        container_name: str,
        connection_string: str,
        blob: Optional[str] = None,
        name_starts_with: Optional[str] = None,
        include: Optional[Any] = None,
        **kwargs: Any,
    ) -> None:
        """Initialize Azure Storage Blob reader.
        
        Args:
            container_name: Name of the Azure container
            connection_string: Azure storage connection string
            blob: Specific blob name to download (optional)
            name_starts_with: Filter blobs by name prefix (optional)
            include: Additional datasets to include in response (optional)
        """
        super().__init__(**kwargs)
        self.container_name = container_name
        self.connection_string = connection_string
        self.blob = blob
        self.name_starts_with = name_starts_with
        self.include = include
        
        # Custom file extractors that don't chunk
        self.file_extractors = {
            ".pdf": SimplePDFReader(),
            ".docx": SimpleDocxReader(),
        }
    
    def load_data(self) -> List[Document]:
        """Load documents from Azure Storage Blob."""
        try:
            from azure.storage.blob import ContainerClient
        except ImportError as exc:
            raise ImportError(
                "Could not import azure storage blob python package. "
                "Please install it with `pip install azure-storage-blob`."
            ) from exc
        
        container_client = ContainerClient.from_connection_string(
            conn_str=self.connection_string, 
            container_name=self.container_name
        )
        
        documents = []
        
        with tempfile.TemporaryDirectory() as temp_dir:
            downloaded_files = []
            
            if self.blob:
                # Download specific blob
                downloaded_files = [self._download_blob(container_client, self.blob, temp_dir)]
            else:
                # Download all blobs matching criteria
                print("Listing blobs...")
                blobs_list = container_client.list_blobs(
                    name_starts_with=self.name_starts_with, 
                    include=self.include
                )
                
                for blob_obj in blobs_list:
                    file_path = self._download_blob(container_client, blob_obj.name, temp_dir)
                    downloaded_files.append(file_path)
            
            # Process downloaded files
            for file_path in downloaded_files:
                if file_path:
                    docs = self._process_file(file_path)
                    documents.extend(docs)
        
        print(f"Loaded {len(documents)} documents (no chunking applied)")
        return documents
    
    def _download_blob(self, container_client, blob_name: str, temp_dir: str) -> Optional[str]:
        """Download a single blob and return the file path."""
        try:
            extension = Path(blob_name).suffix
            download_file_path = f"{temp_dir}/{Path(blob_name).stem}{extension}"
            
            print(f"Downloading {blob_name}...")
            start_time = time.time()
            
            stream = container_client.download_blob(blob_name)
            with open(download_file_path, "wb") as download_file:
                stream.readinto(download_file)
            
            end_time = time.time()
            print(f"{blob_name} downloaded in {end_time - start_time:.2f} seconds")
            
            return download_file_path
            
        except Exception as e:
            print(f"Error downloading {blob_name}: {str(e)}")
            return None
    
    def _process_file(self, file_path: str) -> List[Document]:
        """Process a single file using appropriate reader."""
        file_path_obj = Path(file_path)
        extension = file_path_obj.suffix.lower()
        
        # Get appropriate reader for file type
        if extension in self.file_extractors:
            reader = self.file_extractors[extension]
            try:
                # Add blob-specific metadata
                extra_info = {
                    "source": "azure_blob",
                    "container": self.container_name,
                    "file_extension": extension
                }
                
                documents = reader.load_data(file_path_obj, extra_info=extra_info)
                print(f"Processed {file_path_obj.name}: {len(documents)} document(s)")
                return documents
                
            except Exception as e:
                print(f"Error processing {file_path_obj.name}: {str(e)}")
                return []
        else:
            print(f"Unsupported file type: {extension}")
            return []
